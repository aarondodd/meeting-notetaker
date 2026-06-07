"""Markdown -> .docx export with Word's native TOC field (#94).

Two surfaces:

  * ``export_to_docx`` -- pure-Python via ``python-docx``. Renders the
    markdown body into a Word document with Heading 1 / 2 / 3 styles,
    inline images resolved from a base directory, bold + italic runs,
    bulleted + numbered lists, code blocks, links, and a Word native
    TOC field at the top. Cross-platform; no COM dependency. When the
    user opens the resulting .docx in Word they're prompted to update
    the TOC (right-click -> Update Field) -- one click.

  * ``populate_toc_via_word`` and ``export_to_pdf_via_word`` -- Windows-
    only Word COM helpers. The first opens a .docx, updates the TOC
    field, and saves in place so the TOC is populated when re-opened.
    The second additionally exports the doc as PDF via Word's
    ``ExportAsFixedFormat`` (with ``CreateBookmarks=1`` so the heading
    hierarchy lands in the PDF outline AND the TOC entries are
    clickable hyperlinks). The second is the implementation of the
    "use Word for PDF export" setting toggle (#94 follow-up).

The pure-Python path doesn't require Office to be installed; it's
the everyday "Save as Word..." entry point. The COM path only runs
when the user opts in to "use Word for PDF" and is on Windows with
Word installed; otherwise we fall back to Qt's native PDF backend
(which has its own #94 post-processor for clickable TOC entries).
"""
from __future__ import annotations

import logging
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Optional
from urllib.parse import urlparse

import mistune


log = logging.getLogger(__name__)


# Word's native TOC field code. \o "1-N" = include heading levels 1
# through N; \h = make entries hyperlinks; \z = hide tab leaders in
# Web Layout; \u = use the paragraph's outline level rather than
# style name (matches what Word's own Insert -> TOC produces).
_TOC_FIELD_FMT = 'TOC \\o "1-{depth}" \\h \\z \\u'

# Default body image width when the source markdown doesn't specify
# one. ~6 inches wide (page width minus 1.25" margins on each side).
# python-docx wants EMUs but accepts the Inches() helper.
_DEFAULT_IMAGE_WIDTH_INCHES = 6.0


# ---- public types ---------------------------------------------------------


@dataclass
class WordExportStats:
    """Returned by export_to_docx for caller / test inspection."""
    headings_emitted: int = 0
    paragraphs_emitted: int = 0
    images_embedded: int = 0
    images_missing: int = 0
    toc_inserted: bool = False
    error: Optional[str] = None


# ---- platform / dependency probes -----------------------------------------


def is_python_docx_available() -> bool:
    """True if ``python-docx`` is importable in the current env."""
    try:
        import docx  # noqa: F401, PLC0415
        return True
    except ImportError:
        return False


def is_word_com_available() -> bool:
    """True if we're on Windows AND ``win32com`` can dispatch Word.

    Cheap-fail: only the imports are exercised here, not an actual
    DispatchEx (which would launch Word). The full check happens
    naturally inside ``populate_toc_via_word`` / ``export_to_pdf_via_word``
    when the dispatch itself raises if Word isn't installed.
    """
    if not sys.platform.startswith("win"):
        return False
    try:
        import win32com.client  # noqa: F401, PLC0415
        return True
    except ImportError:
        return False


# ---- markdown -> docx -----------------------------------------------------


_PARSER_OPTIONS = {
    "renderer": "ast",
    "plugins": ["strikethrough", "table", "task_lists", "url"],
}


def export_to_docx(
    markdown_text: str,
    dst: Path,
    *,
    base_dir: Optional[Path] = None,
    title: Optional[str] = None,
    include_toc: bool = False,
    toc_max_depth: int = 3,
    skip_toc_list: bool = True,
) -> WordExportStats:
    """Convert ``markdown_text`` to a Word .docx file at ``dst``.

    ``base_dir`` -- used to resolve relative image refs (typically the
    session directory). Local images are embedded inline; remote
    images (http(s)://) become a placeholder paragraph noting the URL
    so the document doesn't silently lose content.

    ``title`` -- if given, becomes the document title style (above all
    content). When omitted, the markdown's first H1 plays that role.

    ``include_toc`` -- inserts a Word native TOC field above the body.
    ``toc_max_depth`` clamps which heading levels participate.

    ``skip_toc_list`` -- when True (default), drops any pre-existing
    "## Contents" + the markdown bullet list that follows it. This is
    the same dedup pattern the Confluence / Notion exports use:
    if the caller already prepended our markdown TOC list and is now
    asking for a native TOC, we don't want both.

    Returns ``WordExportStats`` with counts + error.
    """
    stats = WordExportStats()
    if not is_python_docx_available():
        stats.error = "python-docx is not installed in this environment"
        return stats

    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Inches, Pt

    base_dir = base_dir or Path.cwd()
    document = Document()
    document.core_properties.title = title or ""

    if include_toc:
        _insert_toc_field(document, toc_max_depth=toc_max_depth)
        stats.toc_inserted = True

    parser = mistune.create_markdown(**_PARSER_OPTIONS)
    nodes = parser(markdown_text or "")
    if skip_toc_list:
        nodes = _drop_toc_block(nodes)

    for node in nodes:
        ntype = node.get("type")
        if ntype == "heading":
            level = int(node.get("attrs", {}).get("level", 1))
            text = _flatten_inline_text(node.get("children", []))
            heading = document.add_heading(text, level=min(level, 9))
            stats.headings_emitted += 1
        elif ntype == "paragraph":
            _emit_paragraph(
                document, node, base_dir=base_dir, stats=stats,
            )
            stats.paragraphs_emitted += 1
        elif ntype == "list":
            ordered = bool(node.get("attrs", {}).get("ordered"))
            _emit_list(
                document, node, ordered=ordered,
                base_dir=base_dir, stats=stats,
            )
        elif ntype == "block_code":
            _emit_code_block(document, node)
        elif ntype == "block_quote":
            _emit_quote(document, node)
        elif ntype == "thematic_break":
            document.add_paragraph("_" * 40)
        elif ntype == "block_html":
            # Best-effort: emit raw HTML as a plain paragraph. python-
            # docx can't render arbitrary HTML and we don't want to.
            text = (node.get("raw") or "").strip()
            if text:
                document.add_paragraph(text)
        # Other node types (table, etc.) are not yet supported -- they
        # land as nothing rather than crashing. Tables can be added
        # later if user feedback shows the need.

    try:
        document.save(str(dst))
    except OSError as exc:
        stats.error = str(exc)
        return stats
    return stats


def _insert_toc_field(document, *, toc_max_depth: int) -> None:
    """Append Word's native TOC field to the document.

    The field's instrText is what Word evaluates when it populates
    the TOC. Until Word evaluates it (either via the user clicking
    Update Field, or via our COM step in populate_toc_via_word), the
    placeholder text in the 'separate' run is what shows up.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    r_elem = run._element

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r_elem.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = _TOC_FIELD_FMT.format(depth=max(1, min(toc_max_depth, 9)))
    r_elem.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r_elem.append(fld_sep)

    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "Right-click here and choose 'Update Field' to populate the "
        "table of contents."
    )
    r_elem.append(placeholder)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_elem.append(fld_end)

    # Page break after the TOC so the body content starts on page 2,
    # matching the layout Word produces when you Insert -> TOC by
    # hand. The empty paragraph holding the break is harmless.
    break_para = document.add_paragraph()
    break_run = break_para.add_run()
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    break_run._element.append(page_break)


def _drop_toc_block(nodes: list[dict]) -> list[dict]:
    """Remove a "## Contents" heading and the bullet list right after.

    The synthesis pipeline can prepend a markdown TOC list when the
    target doesn't have a native primitive. Word does, so we strip
    that list to avoid showing both side by side.
    """
    out: list[dict] = []
    # State: 0 = passthrough; 1 = saw "## Contents", waiting for the
    # bullet list; 2 = list consumed, drop one optional <hr/> next.
    # blank_line nodes are skipped transparently in states 1 + 2 so a
    # blank line between Contents / list / hr doesn't break the run.
    state = 0
    for node in nodes:
        ntype = node.get("type")
        if state == 0:
            if ntype == "heading":
                text = _flatten_inline_text(node.get("children", []))
                if text.strip().lower() == "contents":
                    state = 1
                    continue
            out.append(node)
            continue
        if state == 1:
            if ntype == "blank_line":
                continue
            if ntype == "list":
                state = 2
                continue
            # No list after "Contents" -- bail out, restore the node.
            state = 0
            out.append(node)
            continue
        if state == 2:
            if ntype == "blank_line":
                continue
            if ntype == "thematic_break":
                state = 0
                continue
            state = 0
            out.append(node)
    return out


# ---- inline + block helpers -----------------------------------------------


def _flatten_inline_text(children: list[dict]) -> str:
    """Collapse an AST inline-children list into a plain string.

    Used by code paths that need just the text content (e.g. heading
    text for ``add_heading``). Bold / italic styling is lost; that's
    fine for headings where Word's heading style governs appearance.
    """
    parts: list[str] = []
    for child in children or []:
        ctype = child.get("type")
        if ctype == "text":
            parts.append(child.get("raw") or "")
        elif ctype == "codespan":
            parts.append(child.get("raw") or "")
        elif ctype == "linebreak":
            parts.append("\n")
        elif ctype == "softbreak":
            parts.append(" ")
        elif ctype == "link":
            parts.append(_flatten_inline_text(child.get("children", [])))
        elif ctype in ("emphasis", "strong", "strikethrough"):
            parts.append(_flatten_inline_text(child.get("children", [])))
        elif ctype == "image":
            alt = (child.get("attrs") or {}).get("alt") or ""
            parts.append(alt)
    return "".join(parts)


def _emit_paragraph(
    document, node: dict, *, base_dir: Path, stats: WordExportStats,
) -> None:
    """Append a regular markdown paragraph to ``document``.

    Inline images are emitted inline if local + resolvable; remote
    images become a placeholder run noting the URL.
    """
    paragraph = document.add_paragraph()
    _emit_inline_children(
        paragraph, node.get("children", []),
        base_dir=base_dir, stats=stats,
    )


def _emit_list(
    document, node: dict, *, ordered: bool, base_dir: Path,
    stats: WordExportStats,
) -> None:
    """Append a bulleted or numbered list. Nested lists handled
    recursively; the inner list's items use Word's level-2 list style.
    """
    style = "List Number" if ordered else "List Bullet"
    for item in node.get("children", []):
        # Each list item is a block_text + optional nested list.
        for child in item.get("children", []):
            if child.get("type") == "block_text":
                paragraph = document.add_paragraph(style=style)
                _emit_inline_children(
                    paragraph, child.get("children", []),
                    base_dir=base_dir, stats=stats,
                )
            elif child.get("type") == "paragraph":
                paragraph = document.add_paragraph(style=style)
                _emit_inline_children(
                    paragraph, child.get("children", []),
                    base_dir=base_dir, stats=stats,
                )
            elif child.get("type") == "list":
                # Nested list -- recurse with a sub-style.
                inner_ordered = bool(
                    child.get("attrs", {}).get("ordered"),
                )
                inner_style = (
                    "List Number 2" if inner_ordered else "List Bullet 2"
                )
                for inner_item in child.get("children", []):
                    for inner_child in inner_item.get("children", []):
                        if inner_child.get("type") in (
                            "block_text", "paragraph",
                        ):
                            inner_p = document.add_paragraph(
                                style=inner_style,
                            )
                            _emit_inline_children(
                                inner_p,
                                inner_child.get("children", []),
                                base_dir=base_dir, stats=stats,
                            )


def _emit_code_block(document, node: dict) -> None:
    """Append a fenced code block as a monospaced paragraph."""
    from docx.shared import Pt

    text = node.get("raw") or ""
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["No Spacing"]
    run = paragraph.add_run(text.rstrip("\n"))
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def _emit_quote(document, node: dict) -> None:
    """Append a block quote as a single italic paragraph."""
    paragraph = document.add_paragraph()
    try:
        paragraph.style = document.styles["Intense Quote"]
    except KeyError:
        paragraph.style = document.styles["Normal"]
    for child in node.get("children", []):
        if child.get("type") == "paragraph":
            run = paragraph.add_run(
                _flatten_inline_text(child.get("children", [])),
            )
            run.italic = True


def _emit_inline_children(
    paragraph, children: list[dict], *, base_dir: Path,
    stats: WordExportStats,
) -> None:
    """Append inline content (text + emphasis + links + images) to
    an existing paragraph. The recursive dispatch matches mistune's
    AST: emphasis/strong wrap children, text/codespan/linebreak are
    leaves, image is a leaf, link is an inline wrapper.
    """
    from docx.shared import Inches

    for child in children or []:
        ctype = child.get("type")
        if ctype == "text":
            paragraph.add_run(child.get("raw") or "")
        elif ctype == "codespan":
            run = paragraph.add_run(child.get("raw") or "")
            run.font.name = "Consolas"
        elif ctype == "linebreak":
            paragraph.add_run().add_break()
        elif ctype == "softbreak":
            paragraph.add_run(" ")
        elif ctype == "strong":
            text = _flatten_inline_text(child.get("children", []))
            run = paragraph.add_run(text)
            run.bold = True
        elif ctype == "emphasis":
            text = _flatten_inline_text(child.get("children", []))
            run = paragraph.add_run(text)
            run.italic = True
        elif ctype == "strikethrough":
            text = _flatten_inline_text(child.get("children", []))
            run = paragraph.add_run(text)
            run.font.strike = True
        elif ctype == "link":
            text = _flatten_inline_text(child.get("children", []))
            url = (child.get("attrs") or {}).get("url") or ""
            _add_hyperlink(paragraph, text or url, url)
        elif ctype == "image":
            url = (child.get("attrs") or {}).get("url") or ""
            alt = (child.get("attrs") or {}).get("alt") or ""
            if not url:
                continue
            local = _resolve_local_image(url, base_dir=base_dir)
            if local is not None and local.is_file():
                try:
                    run = paragraph.add_run()
                    run.add_picture(
                        str(local),
                        width=Inches(_DEFAULT_IMAGE_WIDTH_INCHES),
                    )
                    stats.images_embedded += 1
                except Exception as exc:  # noqa: BLE001
                    log.warning("docx image embed failed for %s: %s", local, exc)
                    paragraph.add_run(f"(image: {alt or url})")
                    stats.images_missing += 1
            else:
                paragraph.add_run(f"(image: {alt or url})")
                stats.images_missing += 1


_REMOTE_SCHEMES = {"http", "https", "data"}


def _resolve_local_image(url: str, *, base_dir: Path) -> Optional[Path]:
    """Resolve an image url against ``base_dir`` for embedding.

    Returns the path if it's local + present on disk. Remote and
    data: URLs return None (caller emits a placeholder).
    """
    if not url:
        return None
    parsed = urlparse(url)
    if parsed.scheme in _REMOTE_SCHEMES:
        return None
    if parsed.scheme == "file":
        candidate = Path(parsed.path)
    else:
        candidate = (base_dir / url).resolve()
    return candidate if candidate.is_file() else None


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Emit a Word-native hyperlink run.

    python-docx's high-level API doesn't expose hyperlinks; the
    standard workaround is to add a relationship and build the
    ``<w:hyperlink>`` element by hand. The link click in Word jumps
    to the target URL.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    new_run.append(rpr)

    t_elem = OxmlElement("w:t")
    t_elem.text = text or url
    new_run.append(t_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ---- Word COM (Windows-only) ---------------------------------------------


# wdExportFormatPDF, OptimizeForPrint, AllDocument constants. Mirror
# what office_preview._word_to_pdf uses so behaviour is consistent.
_WD_EXPORT_FORMAT_PDF = 17
_WD_EXPORT_OPTIMIZE_FOR_PRINT = 0
_WD_EXPORT_ALL_DOCUMENT = 0
# wdExportCreateHeadingBookmarks = 1 -- emits a PDF outline (sidebar
# bookmarks) for every heading, AND makes the TOC field's entries
# clickable hyperlinks in the resulting PDF. Without this, Word
# writes a PDF with neither.
_WD_EXPORT_CREATE_HEADING_BOOKMARKS = 1


def populate_toc_via_word(
    docx_path: Path, *, save_in_place: bool = True,
) -> bool:
    """Open ``docx_path`` in Word, update its TOC field(s), save.

    Returns True on success, False on any failure (Word not
    available, no TOC field present, etc.). Non-destructive: if
    ``save_in_place`` is False, the document is left modified in
    Word's memory and immediately closed without saving.
    """
    if not is_word_com_available():
        return False
    try:
        import pythoncom  # noqa: PLC0415
        import win32com.client  # noqa: PLC0415
    except ImportError:
        return False

    pythoncom.CoInitialize()
    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(
            FileName=str(docx_path.resolve()),
            ReadOnly=False,
            AddToRecentFiles=False,
            ConfirmConversions=False,
        )
        try:
            tocs = doc.TablesOfContents
            if tocs.Count == 0:
                return False
            for i in range(1, tocs.Count + 1):
                tocs(i).Update()
            if save_in_place:
                doc.Save()
        finally:
            doc.Close(SaveChanges=False)
        return True
    except Exception:  # noqa: BLE001
        log.exception("populate_toc_via_word failed for %s", docx_path)
        return False
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:  # noqa: BLE001
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:  # noqa: BLE001
            pass


def export_to_pdf_via_word(
    docx_path: Path, pdf_path: Path,
) -> bool:
    """Open ``docx_path`` in Word, update TOC fields, export to PDF.

    The PDF carries Word's native outline (sidebar bookmarks per
    heading) AND clickable TOC entries -- both via the
    ``CreateBookmarks=1`` flag on ExportAsFixedFormat. This is the
    cleanest way to get a Word-quality PDF with working navigation.

    Returns True on success, False on failure. The .docx is left on
    disk unchanged regardless of success (Word's edits are not
    saved back unless the caller previously called
    ``populate_toc_via_word(..., save_in_place=True)``).
    """
    if not is_word_com_available():
        return False
    try:
        import pythoncom  # noqa: PLC0415
        import win32com.client  # noqa: PLC0415
    except ImportError:
        return False

    pythoncom.CoInitialize()
    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(
            FileName=str(docx_path.resolve()),
            ReadOnly=False,
            AddToRecentFiles=False,
            ConfirmConversions=False,
        )
        try:
            try:
                tocs = doc.TablesOfContents
                for i in range(1, tocs.Count + 1):
                    tocs(i).Update()
            except Exception:  # noqa: BLE001
                # No TOC field is not fatal; we still want the PDF.
                pass
            doc.ExportAsFixedFormat(
                OutputFileName=str(pdf_path.resolve()),
                ExportFormat=_WD_EXPORT_FORMAT_PDF,
                OpenAfterExport=False,
                OptimizeFor=_WD_EXPORT_OPTIMIZE_FOR_PRINT,
                Range=_WD_EXPORT_ALL_DOCUMENT,
                CreateBookmarks=_WD_EXPORT_CREATE_HEADING_BOOKMARKS,
                DocStructureTags=False,
                BitmapMissingFonts=True,
                UseISO19005_1=False,
            )
        finally:
            doc.Close(SaveChanges=False)
        return True
    except Exception:  # noqa: BLE001
        log.exception("export_to_pdf_via_word failed for %s", docx_path)
        return False
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:  # noqa: BLE001
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:  # noqa: BLE001
            pass
